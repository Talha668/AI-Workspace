import PyPDF2
import docx
import io
import magic
from django.core.files.uploadedfile import UploadedFile
from .models import Document


class DocumentProcessor:
    @staticmethod
    def detect_file_type(file):
        """Detect the MIME type and return our document type"""
        mime = magic.from_buffer(file.read(1024), mime=True)
        file.seek(0)  # Reset file pointer
        
        mime_to_type = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain': 'txt',
        }
        return mime_to_type.get(mime, 'txt')

    @staticmethod
    def extract_text(file_field, file_type):
        """Extract text from different document types"""
        try:
            if file_type == 'pdf':
                return DocumentProcessor._extract_pdf(file_field)
            elif file_type == 'docx':
                return DocumentProcessor._extract_docx(file_field)
            elif file_type == 'txt':
                return DocumentProcessor._extract_txt(file_field)
            return ''
        except Exception as e:
            print(f"Error extracting text: {str(e)}")
            return ''

    @staticmethod
    def _extract_pdf(file_field):
        text = ''
        pdf_reader = PyPDF2.PdfReader(file_field)
        for page in pdf_reader.pages:
            text += page.extract_text() + '\n'
        return text

    @staticmethod
    def _extract_docx(file_field):
        text = ''
        doc = docx.Document(file_field)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
        return text

    @staticmethod
    def _extract_txt(file_field):
        return file_field.read().decode('utf-8')

class DocumentService:
    @staticmethod
    def process_document(file: UploadedFile, user, workspace):
        """Process an uploaded document"""
        # Detect file type
        file_type = DocumentProcessor.detect_file_type(file)
        
        # Extract text
        content_text = DocumentProcessor.extract_text(file, file_type)
        
        # Create document record
        document = Document(
            workspace=workspace,
            title=file.name,
            file=file,
            file_type=file_type,
            file_size=file.size,
            content_text=content_text,
            uploaded_by=user,
            is_processed=True
        )
        document.save()
        
        return document